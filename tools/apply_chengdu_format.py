"""
成都大学本科毕业论文格式套用
Pipeline C: 根据学校规范格式化论文
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, font_name_cn='宋体', font_name_en='Times New Roman', font_size=12):
    """设置中英文字体和大小"""
    run.font.name = font_name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    run.font.size = Pt(font_size)

def set_paragraph_format(para, font_size=12, first_line_indent=True, line_spacing=1.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """设置段落格式"""
    para.alignment = alignment
    
    # 设置行距
    para.paragraph_format.line_spacing = line_spacing
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    
    # 首行缩进
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(0.74)  # 两个字符
    
    # 段后间距
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)

def is_heading(text, level=1):
    """识别是否为指定级别的标题"""
    # 一级标题：第X章 + 标题
    if level == 1:
        return bool(re.match(r'^第[一二三四五六七八九十]+章\s*[\u4e00-\u9fa5]', text))
    # 二级标题：X.Y 或 X.Y.Z 格式
    elif level == 2:
        return bool(re.match(r'^\d+\.\d+[\u4e00-\u9fa5]', text))
    # 三级标题：(X) 或 X、格式
    elif level == 3:
        return bool(re.match(r'^[（(]\d+[）)]\s*[\u4e00-\u9fa5]', text))
    return False

def identify_paragraph_type(text, prev_type=None):
    """识别段落类型"""
    text = text.strip()
    if not text:
        return 'empty'
    
    # 封面信息
    if any(kw in text for kw in ['题 目', '学 院', '专 业', '学 生', '学 号', '指 导', '完 成', '班 级']):
        return 'cover'
    
    # 声明类
    if '声明' in text or '授权' in text:
        return 'declaration'
    
    # 摘要标题
    if text == '摘要' or text == 'Abstract':
        return 'abstract_title'
    
    # 关键词
    if '关键词' in text or text.startswith('关键词：') or 'Keyword' in text:
        return 'keywords'
    
    # 一级标题（第X章）
    if re.match(r'^第[一二三四五六七八九十]+章', text):
        return 'heading1'
    
    # 二级标题（1.2 格式）
    if re.match(r'^\d+\.\d+', text) and len(text) < 50:
        return 'heading2'
    
    # 三级标题
    if re.match(r'^[（(]\d+[）)]', text) and len(text) < 50:
        return 'heading3'
    
    # 特殊章节标题
    if text in ['绪论', '结论', '参考文献', '致谢', '附录', 'Abstract']:
        return 'heading1'
    
    # 图片说明
    if text.startswith('图') and '示' in text and len(text) < 50:
        return 'figure_caption'
    
    # 表格标题
    if text.startswith('表') and len(text) < 50:
        return 'table_caption'
    
    # 参考文献条目
    if re.match(r'^\d+\.', text) and prev_type == 'reference':
        return 'reference'
    
    # 普通正文
    return 'body'

def format_thesis(input_path, output_path):
    """格式化论文"""
    print(f"=== 读取源文档: {input_path} ===")
    doc = Document(input_path)
    
    # 创建新文档
    new_doc = Document()
    
    # 复制样式
    for style in doc.styles:
        try:
            new_doc.styles.add_style(style.name, style.type)
        except:
            pass
    
    # 设置页面格式 (A4)
    section = new_doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    
    # 统计
    stats = {'cover': 0, 'heading1': 0, 'heading2': 0, 'heading3': 0, 'body': 0, 'empty': 0}
    
    # 处理每个段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        prev_type = stats.get('last_type', 'body')
        para_type = identify_paragraph_type(text, prev_type)
        stats[para_type] = stats.get(para_type, 0) + 1
        stats['last_type'] = para_type
        
        if para_type == 'empty':
            new_doc.add_paragraph()
            continue
        
        # 创建新段落
        new_para = new_doc.add_paragraph()
        
        # 根据类型设置格式
        if para_type == 'cover':
            # 封面信息 - 宋体三号加粗居中
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_para.add_run(text)
            set_font(run, '宋体', 'Times New Roman', 16)
            stats['cover'] += 1
            
        elif para_type == 'declaration':
            # 声明 - 宋体小四
            new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = new_para.add_run(text)
            set_font(run, '宋体', 'Times New Roman', 12)
            new_para.paragraph_format.first_line_indent = Cm(0.74)
            new_para.paragraph_format.line_spacing = 1.5
            
        elif para_type == 'abstract_title':
            # 摘要标题 - 黑体三号居中
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_para.add_run(text)
            set_font(run, '黑体', 'Times New Roman', 16)
            
        elif para_type == 'keywords':
            # 关键词 - 宋体小四
            new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = new_para.add_run(text)
            set_font(run, '宋体', 'Times New Roman', 12)
            
        elif para_type == 'heading1':
            # 一级标题 - 黑体三号居中
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_para.add_run(text.replace('chapter', ''))
            set_font(run, '黑体', 'Times New Roman', 18)
            # 段前段后间距
            new_para.paragraph_format.space_before = Pt(24)
            new_para.paragraph_format.space_after = Pt(18)
            stats['heading1'] += 1
            
        elif para_type == 'heading2':
            # 二级标题 - 黑体四号
            new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = new_para.add_run(text)
            set_font(run, '黑体', 'Times New Roman', 14)
            new_para.paragraph_format.space_before = Pt(18)
            new_para.paragraph_format.space_after = Pt(12)
            stats['heading2'] += 1
            
        elif para_type == 'heading3':
            # 三级标题 - 黑体小四
            new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = new_para.add_run(text)
            set_font(run, '黑体', 'Times New Roman', 12)
            new_para.paragraph_format.space_before = Pt(12)
            new_para.paragraph_format.space_after = Pt(6)
            stats['heading3'] += 1
            
        elif para_type in ['figure_caption', 'table_caption']:
            # 图片/表格标题 - 宋体五号居中
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_para.add_run(text)
            set_font(run, '宋体', 'Times New Roman', 10.5)
            
        elif para_type == 'reference':
            # 参考文献 - 宋体小四
            new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = new_para.add_run(text)
            set_font(run, '宋体', 'Times New Roman', 12)
            new_para.paragraph_format.first_line_indent = Cm(-0.74)
            new_para.paragraph_format.left_indent = Cm(0.74)
            new_para.paragraph_format.line_spacing = 1.5
            
        else:  # body
            # 正文 - 宋体小四，首行缩进，1.5倍行距
            new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = new_para.add_run(text)
            set_font(run, '宋体', 'Times New Roman', 12)
            new_para.paragraph_format.first_line_indent = Cm(0.74)
            new_para.paragraph_format.line_spacing = 1.5
            new_para.paragraph_format.space_after = Pt(0)
            new_para.paragraph_format.space_before = Pt(0)
            stats['body'] += 1
    
    # 复制表格
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text
                # 表格内容居中
                for p in new_table.rows[i].cells[j].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        set_font(run, '宋体', 'Times New Roman', 10.5)
        new_doc.add_paragraph()  # 表格后空行
    
    # 保存
    new_doc.save(output_path)
    print(f"\n=== 保存到: {output_path} ===")
    
    # 打印统计
    print(f"\n=== 格式处理统计 ===")
    print(f"  封面信息: {stats['cover']}")
    print(f"  一级标题: {stats['heading1']}")
    print(f"  二级标题: {stats['heading2']}")
    print(f"  三级标题: {stats['heading3']}")
    print(f"  正文段落: {stats['body']}")
    print(f"  表格数: {len(doc.tables)}")
    
    return output_path

if __name__ == '__main__':
    input_file = 'thesis_yoga.docx'
    output_dir = 'outputs'
    output_file = os.path.join(output_dir, '成都大学_瑜伽动作评估系统_格式化.docx')
    
    os.makedirs(output_dir, exist_ok=True)
    
    print("=" * 60)
    print("成都大学本科毕业论文格式套用")
    print("=" * 60)
    
    if os.path.exists(input_file):
        format_thesis(input_file, output_file)
        print("\n[OK] 完成！")
        print(f"\n输出文件: {output_file}")
    else:
        print(f"错误: 找不到文件 {input_file}")
