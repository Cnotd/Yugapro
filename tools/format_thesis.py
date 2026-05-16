"""
成都大学本科毕业论文格式化
直接应用学校标准的论文格式
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def apply_chengdu_format(input_path, output_path):
    """应用成都大学本科毕业论文格式"""
    print(f"读取: {input_path}")
    doc = Document(input_path)
    
    # 创建新文档
    new_doc = Document()
    
    # === 页面设置 ===
    section = new_doc.sections[0]
    section.page_width = Cm(21)      # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    
    def set_run_font(run, cn='宋体', en='Times New Roman', size=12, bold=False):
        """设置字体"""
        run.font.name = en
        run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)
        run.font.size = Pt(size)
        run.bold = bold
    
    def add_para(text, align='justify', cn='宋体', size=12, bold=False, 
                  indent=True, spacing=1.5, before=0, after=0):
        """添加格式化段落"""
        para = new_doc.add_paragraph()
        para.alignment = {
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'left': WD_ALIGN_PARAGRAPH.LEFT
        }.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
        
        para.paragraph_format.line_spacing = spacing
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after = Pt(after)
        
        if indent:
            para.paragraph_format.first_line_indent = Cm(0.74)
        
        if text:
            run = para.add_run(text)
            set_run_font(run, cn, 'Times New Roman', size, bold)
        
        return para
    
    # === 识别并处理每个段落 ===
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 空行
        if not text:
            add_para('', indent=False, spacing=1.0)
            continue
        
        # === 封面信息 (居中) ===
        if any(kw in text for kw in ['本科毕业设计', '题 目', '学  院', '专 业', 
                                      '学 生', '学 号', '指 导', '完 成', '班 级']):
            add_para(text, 'center', '宋体', 16, True)
            continue
        
        # === 声明类 ===
        if '声明' in text or '授权' in text:
            add_para(text, 'justify', '宋体', 12)
            continue
        
        # === 签名栏 ===
        if '论文作者签名' in text or '指导教师签名' in text:
            add_para(text, 'center', '宋体', 12)
            continue
        
        # === 摘要标题 ===
        if text == '摘要':
            add_para(text, 'center', '黑体', 16, True, before=24, after=18)
            continue
        if text == 'Abstract':
            add_para(text, 'center', '黑体', 16, True, before=12, after=18)
            continue
        
        # === 关键词 ===
        if '关键词' in text or 'Keyword' in text.upper():
            add_para(text, 'left', '黑体', 12, True, indent=False)
            continue
        
        # === 一级标题 (第X章 / 绪论 / 结论等) ===
        if re.match(r'^第[一二三四五六七八九十]+章', text) or \
           text in ['绪论', '结论', '参考文献', '致谢', '附录']:
            text = text.replace('chapter', '').strip()
            add_para(text, 'center', '黑体', 18, True, indent=False, before=24, after=18)
            continue
        
        # === 二级标题 (X.X 格式) ===
        if re.match(r'^\d+\.\d+\s+[\u4e00-\u9fa5]', text):
            add_para(text, 'left', '黑体', 14, True, indent=False, before=18, after=12)
            continue
        
        # === 三级标题 ((X) 格式) ===
        if re.match(r'^[（(]\d+[）)]\s*[\u4e00-\u9fa5]', text):
            add_para(text, 'left', '黑体', 12, True, indent=False, before=12, after=6)
            continue
        
        # === 摘要正文 ===
        if text.startswith('本课题') or text.startswith('本文') or \
           text.startswith('This') or text.startswith('This thesis'):
            add_para(text, 'justify', '宋体', 12)
            continue
        
        # === 正文 ===
        add_para(text, 'justify', '宋体', 12)
    
    # === 复制表格 ===
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text
                for p in new_table.rows[i].cells[j].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        set_run_font(run, '宋体', 'Times New Roman', 10.5)
        add_para('', indent=False)  # 表格后空行
    
    # === 保存 ===
    new_doc.save(output_path)
    print(f"保存: {output_path}")

if __name__ == '__main__':
    input_file = 'thesis_yoga.docx'
    output_file = 'outputs/成都大学_瑜伽动作评估系统_论文.docx'
    
    os.makedirs('outputs', exist_ok=True)
    
    print("=" * 60)
    print("成都大学本科毕业论文格式套用")
    print("=" * 60)
    
    if os.path.exists(input_file):
        apply_chengdu_format(input_file, output_file)
        print("\n" + "=" * 60)
        print("完成！请打开 outputs 目录查看格式化后的文档")
        print("=" * 60)
    else:
        print(f"错误: 找不到 {input_file}")
