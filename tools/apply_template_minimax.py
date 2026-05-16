# -*- coding: utf-8 -*-
"""
成都大学毕业论文 - 套用学校模板
使用 Python 和 python-docx 实现 minimax-docx 技能的工作流程

Pipeline C: Apply Template
参考: minimax-docx/references/scenario_c_apply_template.md
"""

import os
import sys
import shutil
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# 路径配置
BASE_DIR = Path(r"d:\yuga_test")
SOURCE = BASE_DIR / "thesis_yoga.docx"
TEMPLATE = BASE_DIR / "成都大学本科毕业设计（论文）模板.docx"
OUTPUT = BASE_DIR / "outputs" / "thesis_chengdu_university.docx"
LOGS_DIR = BASE_DIR / "logs"
TEMP_DIR = BASE_DIR / "temp"

def setup_directories():
    """创建必要目录"""
    for d in [OUTPUT.parent, LOGS_DIR, TEMP_DIR]:
        d.mkdir(parents=True, exist_ok=True)
    print(f"[设置] 输出目录: {OUTPUT.parent}")

def analyze_document(doc_path, name="文档"):
    """分析文档结构和样式"""
    print(f"\n[分析] {name}: {doc_path}")
    doc = Document(doc_path)
    
    info = {
        "total_paragraphs": 0,
        "total_tables": 0,
        "styles_used": set(),
        "heading_count": 0,
        "first_lines": [],
    }
    
    # 统计段落
    for para in doc.paragraphs:
        info["total_paragraphs"] += 1
        if para.style:
            info["styles_used"].add(para.style.name)
            if "Heading" in para.style.name or "标题" in para.style.name:
                info["heading_count"] += 1
        if para.text.strip() and len(info["first_lines"]) < 10:
            info["first_lines"].append(para.text[:80])
    
    # 统计表格
    info["total_tables"] = len(doc.tables)
    
    print(f"  - 段落数: {info['total_paragraphs']}")
    print(f"  - 表格数: {info['total_tables']}")
    print(f"  - 标题数: {info['heading_count']}")
    print(f"  - 样式: {', '.join(list(info['styles_used'])[:10])}")
    
    # 保存分析结果
    analysis_file = LOGS_DIR / f"{name}_analysis.json"
    with open(analysis_file, 'w', encoding='utf-8') as f:
        json.dump({**info, "styles_used": list(info["styles_used"])}, f, ensure_ascii=False, indent=2)
    print(f"  - 分析结果: {analysis_file}")
    
    return info

def extract_content(source_path):
    """从源文档提取内容"""
    print(f"\n[提取] 从源文档提取内容...")
    source = Document(source_path)
    
    content_items = []
    
    for para in source.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        item = {
            "text": text,
            "style": para.style.name if para.style else "Normal",
            "alignment": str(para.alignment) if para.alignment else "None",
        }
        
        # 提取格式信息
        if para.runs:
            first_run = para.runs[0]
            item["bold"] = first_run.bold
            item["italic"] = first_run.italic
            item["font_size"] = str(first_run.font.size) if first_run.font.size else None
        
        content_items.append(item)
    
    print(f"  - 提取了 {len(content_items)} 个段落")
    return content_items

def apply_template():
    """
    应用模板 - Pipeline C: Base-Replace 策略
    
    策略说明 (根据 scenario_c_apply_template.md):
    - Zone A (前置页面): 保留模板
    - Zone B (正文内容): 替换为源文档内容
    - Zone C (后置页面): 保留模板
    """
    print("\n[执行] 套用模板...")
    print("  策略: C-2 Base-Replace (以模板为基，替换正文)")
    
    # 加载模板 (作为基础)
    print(f"\n[加载] 模板: {TEMPLATE}")
    template = Document(TEMPLATE)
    
    # 分析模板结构
    template_info = analyze_document(TEMPLATE, "模板")
    
    # 提取源文档内容
    content = extract_content(SOURCE)
    
    # 创建新文档 (复制模板)
    print(f"\n[处理] 创建输出文档...")
    
    # 方法: 直接复制模板，然后清空正文并添加新内容
    output = Document(TEMPLATE)
    
    # 找到需要替换的正文区域
    body = output._element.body
    
    # 遍历 body 中的所有段落，找到正文开始位置
    # 通常在目录之后，参考文献之前
    
    paras_to_remove = []
    in_body_content = False
    body_start_idx = 0
    body_end_idx = len(list(body))
    
    # 简单策略: 保留模板中前 N 个段落 (封面、声明、摘要等)
    # 然后添加源文档内容
    keep_first = 5  # 保留前5个段落 (需要根据实际模板调整)
    
    # 删除模板中的示例正文内容
    for i, child in enumerate(list(body)):
        if child.tag.endswith('}p'):  # 段落
            text = ''.join(t.text for t in child.iter() if t.text)
            if text.strip():
                # 检查是否到达正文开始 (通常是 "第1章" 或 "第一章")
                if "第1章" in text or "第一章" in text or "绪论" in text:
                    in_body_content = True
                    body_start_idx = i
                # 检查是否到达参考文献
                if "参考文献" in text or "致谢" in text:
                    body_end_idx = i
                    break
    
    print(f"  - 正文开始索引: {body_start_idx}")
    print(f"  - 正文结束索引: {body_end_idx}")
    
    # 收集需要删除的段落
    children_to_remove = list(body)[body_start_idx:body_end_idx]
    for child in children_to_remove:
        body.remove(child)
    
    print(f"  - 删除了 {len(children_to_remove)} 个模板示例段落")
    
    # 添加源文档内容
    for item in content:
        # 创建新段落
        new_para = output.add_paragraph()
        new_para.text = item["text"]
        
        # 尝试应用样式
        try:
            if "Heading1" in item["style"] or "1级" in item["style"]:
                new_para.style = "Heading 1"
            elif "Heading2" in item["style"] or "2级" in item["style"]:
                new_para.style = "Heading 2"
            elif "Heading3" in item["style"] or "3级" in item["style"]:
                new_para.style = "Heading 3"
        except:
            pass
    
    print(f"  - 添加了 {len(content)} 个新段落")
    
    # 保存输出
    print(f"\n[保存] 输出文档: {OUTPUT}")
    output.save(str(OUTPUT))
    
    return True

def validate_output():
    """验证输出文档"""
    print("\n[验证] 检查输出文档...")
    
    if not OUTPUT.exists():
        print("  [错误] 输出文件不存在!")
        return False
    
    size = OUTPUT.stat().st_size
    print(f"  - 文件大小: {size / 1024:.2f} KB")
    
    try:
        doc = Document(OUTPUT)
        para_count = len([p for p in doc.paragraphs if p.text.strip()])
        table_count = len(doc.tables)
        print(f"  - 段落数: {para_count}")
        print(f"  - 表格数: {table_count}")
        print("  [通过] 输出文档有效")
        return True
    except Exception as e:
        print(f"  [错误] {e}")
        return False

def main():
    """主函数"""
    print("=" * 60)
    print("  成都大学毕业论文 - 套用学校模板")
    print("  minimax-docx Skill Pipeline C (Python版)")
    print("=" * 60)
    
    # 检查文件
    if not SOURCE.exists():
        print(f"[错误] 找不到源文档: {SOURCE}")
        return False
    
    if not TEMPLATE.exists():
        print(f"[错误] 找不到模板: {TEMPLATE}")
        return False
    
    # 执行
    setup_directories()
    
    # 分析
    source_info = analyze_document(SOURCE, "源文档")
    
    # 应用模板
    if apply_template():
        # 验证
        if validate_output():
            print("\n" + "=" * 60)
            print("  [成功] 模板套用完成!")
            print("=" * 60)
            print(f"\n输出文件: {OUTPUT}")
            print("\n后续步骤:")
            print("  1. 用 Word 打开输出文件检查格式")
            print("  2. 补充封面页信息 (学号、姓名、专业、导师等)")
            print("  3. 如有格式问题，在 Word 中手动调整样式")
            
            # 打开文件位置
            os.startfile(OUTPUT.parent)
            return True
    
    return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
