"""
成都大学本科毕业论文格式 - 完全符合撰写规范
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

def set_run_font(run, cn='宋体', size=10.5, bold=False):
    """设置字体"""
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    run.font.size = Pt(size)
    run.bold = bold

def add_para(doc, text='', align='justify', cn='宋体', size=10.5, bold=False, 
             indent=True, space_before=0, space_after=0, line_spacing=1.0):
    """添加格式化段落"""
    para = doc.add_paragraph()
    
    # 对齐方式
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # 行距
    para.paragraph_format.line_spacing = line_spacing
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE if line_spacing != 1.0 else WD_LINE_SPACING.SINGLE
    
    # 段前段后间距
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    
    # 首行缩进 (2字符 ≈ 0.74cm)
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    
    if text:
        run = para.add_run(text)
        set_run_font(run, cn, size, bold)
    
    return para

def identify_type(text, prev_type='body'):
    """识别段落类型"""
    text = text.strip()
    if not text:
        return 'empty'
    
    # 封面
    if any(k in text for k in ['本科毕业设计', '题 目', '学  院', '学 院', '专 业', 
                                  '学 生', '学 号', '指 导', '完 成', '班 级']):
        return 'cover'
    
    # 原创性声明
    if '声明' in text or '授权' in text or '知识产权' in text:
        return 'declaration'
    
    # 签名
    if '论文作者签名' in text or '指导教师签名' in text:
        return 'signature'
    
    # 英文封面信息
    if text.startswith('Yoga') or text.startswith('Major') or text.startswith('Student'):
        return 'en_cover'
    
    # 中文摘要标题
    if text == '摘要':
        return 'abstract_title_cn'
    
    # 英文摘要标题
    if text == 'Abstract':
        return 'abstract_title_en'
    
    # 关键词
    if '关键词' in text and len(text) < 50:
        return 'keywords_cn'
    if text.lower().startswith('keywords') or text.lower().startswith('key words'):
        return 'keywords_en'
    
    # 一级标题：第X章 / 绪论 / 结论 / 参考文献 / 致谢 / 系统章节
    if re.match(r'^第[一二三四五六七八九十]+章', text):
        return 'heading1'
    if text in ['绪论', '结论', '参考文献', '致谢', '附录']:
        return 'heading1'
    # 系统设计章节 (系统需求、系统的设计、系统实现、系统测试等)
    if text in ['系统需求', '系统的设计', '系统设计', '系统实现', '系统测试', '系统架构', '系统总体设计', '系统详细设计']:
        return 'heading1'
    
    # 二级标题：1.X 格式
    if re.match(r'^\d+\.\d+\s+[\u4e00-\u9fa5]', text):
        return 'heading2'
    
    # 三级标题：1.1.1 格式
    if re.match(r'^\d+\.\d+\.\d+\s+[\u4e00-\u9fa5]', text):
        return 'heading3'
    
    # 参考文献条目
    if re.match(r'^\[\d+\]', text):
        return 'reference'
    
    # 正文
    return 'body'

def apply_header_footer(section):
    """添加页眉页脚"""
    # 页眉
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run('成都大学学士学位论文（设计）')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)  # 小五号
    
    # 页脚
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)  # 小五号
    
    # 添加页码字段
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def format_thesis(input_path, output_path):
    """格式化论文"""
    print(f"读取: {input_path}")
    doc = Document(input_path)
    new_doc = Document()
    
    # === 页面设置 ===
    section = new_doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(2.5)
    
    # 添加页眉页脚
    apply_header_footer(section)
    
    # 统计
    stats = {'cover': 0, 'heading1': 0, 'heading2': 0, 'body': 0, 'empty': 0}
    prev_type = 'body'
    
    for para in doc.paragraphs:
        text = para.text.strip()
        para_type = identify_type(text, prev_type)
        prev_type = para_type
        
        # 空行
        if para_type == 'empty':
            add_para(new_doc, '', indent=False)
            continue
        
        # === 封面信息 ===
        if para_type == 'cover':
            stats['cover'] += 1
            add_para(new_doc, text, 'center', '宋体', 18, True, indent=False)  # 小二号=18pt
            
        # === 声明 ===
        elif para_type == 'declaration':
            add_para(new_doc, text, 'justify', '宋体', 10.5, indent=True)
            
        # === 签名 ===
        elif para_type == 'signature':
            add_para(new_doc, text, 'center', '宋体', 10.5)
            
        # === 英文封面 ===
        elif para_type == 'en_cover':
            add_para(new_doc, text, 'center', 'Times New Roman', 12)
            
        # === 摘要标题 ===
        elif para_type == 'abstract_title_cn':
            add_para(new_doc, text, 'center', '黑体', 10.5, True, indent=False, 
                    space_before=24, space_after=12)
            
        elif para_type == 'abstract_title_en':
            add_para(new_doc, text, 'center', '黑体', 10.5, True, indent=False,
                    space_before=12, space_after=12)
            
        # === 关键词 ===
        elif para_type in ['keywords_cn', 'keywords_en']:
            cn = '黑体' if para_type == 'keywords_cn' else 'Times New Roman'
            add_para(new_doc, text, 'left', cn, 10.5, True, indent=False)
            
        # === 一级标题 (第X章/绪论/结论等) ===
        elif para_type == 'heading1':
            stats['heading1'] += 1
            text = text.replace('chapter', '').strip()
            add_para(new_doc, text, 'center', '黑体', 14, True, indent=False,  # 四号=14pt
                    space_before=18, space_after=12)
            
        # === 二级标题 (1.X) ===
        elif para_type == 'heading2':
            stats['heading2'] += 1
            add_para(new_doc, text, 'left', '黑体', 12, True, indent=True,  # 小四=12pt
                    space_before=12, space_after=6)
            
        # === 三级标题 (1.1.1) ===
        elif para_type == 'heading3':
            add_para(new_doc, text, 'left', '黑体', 10.5, True, indent=True,  # 五号=10.5pt
                    space_before=6, space_after=6)
            
        # === 参考文献 ===
        elif para_type == 'reference':
            add_para(new_doc, text, 'justify', '宋体', 9, indent=False)  # 小五号=9pt
            
        # === 正文 ===
        else:
            stats['body'] += 1
            add_para(new_doc, text, 'justify', '宋体', 10.5, indent=True,  # 五号=10.5pt
                    space_before=0, space_after=0, line_spacing=1.0)  # 单倍行距
    
    # 复制表格
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text
                for p in new_table.rows[i].cells[j].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        set_run_font(run, '宋体', 9)  # 表格内容小五号
        add_para(new_doc, '', indent=False)
    
    new_doc.save(output_path)
    print(f"保存: {output_path}")
    
    # 统计
    print(f"\n=== 格式处理统计 ===")
    print(f"  封面信息: {stats['cover']}")
    print(f"  一级标题: {stats['heading1']}")
    print(f"  二级标题: {stats['heading2']}")
    print(f"  正文段落: {stats['body']}")
    
    return output_path

if __name__ == '__main__':
    input_file = 'thesis_yoga.docx'
    output_file = 'outputs/成都大学_瑜伽动作评估系统_论文_规范版.docx'
    
    os.makedirs('outputs', exist_ok=True)
    
    print("=" * 60)
    print("成都大学本科毕业论文格式套用 (完全符合撰写规范)")
    print("=" * 60)
    
    if os.path.exists(input_file):
        format_thesis(input_file, output_file)
        print("\n完成！")
    else:
        print(f"错误: 找不到 {input_file}")
