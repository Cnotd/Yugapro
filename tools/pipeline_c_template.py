"""
Pipeline C: 应用成都大学论文模板
根据 scenario_c_apply_template.md 规范执行
"""
import os
import zipfile
import shutil
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def check_files():
    """Step 1: 检查文件"""
    print("=== Step 1: 检查文件 ===")
    
    template = '成都大学本科毕业设计（论文）模板.doc'
    source = 'thesis_yoga.docx'
    output_dir = 'outputs'
    
    os.makedirs(output_dir, exist_ok=True)
    
    for f in [template, source]:
        if os.path.exists(f):
            size = os.path.getsize(f)
            print(f"  {f}: {size:,} bytes")
            
            # 检查是否是有效的 DOCX
            try:
                with zipfile.ZipFile(f, 'r') as z:
                    has_doc = any('word/document.xml' in n for n in z.namelist())
                    print(f"    有效 DOCX: {has_doc}")
            except zipfile.BadZipFile:
                print(f"    格式: .doc (旧格式 - 需要转换)")
        else:
            print(f"  {f}: 不存在!")
            
    return template, source, output_dir

def convert_doc_to_docx(doc_path, output_path):
    """转换 .doc 为 .docx"""
    print(f"\n=== 转换 .doc → .docx ===")
    print(f"  源文件: {doc_path}")
    
    # 方法1: 使用 LibreOffice (如果有)
    import subprocess
    libreoffice_paths = [
        'C:\\Program Files\\LibreOffice\\program\\soffice.exe',
        'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe',
        'soffice'  # PATH 中
    ]
    
    for lo in libreoffice_paths:
        try:
            result = subprocess.run(
                [lo, '--headless', '--convert-to', 'docx', '--outdir', 
                 os.path.dirname(output_path), doc_path],
                capture_output=True, timeout=30
            )
            if result.returncode == 0:
                print(f"  使用 LibreOffice 转换成功")
                return True
        except:
            continue
    
    print("  LibreOffice 不可用，尝试其他方法...")
    return False

def analyze_template(template_path):
    """Step 2: 分析模板结构"""
    print("\n=== Step 2: 分析模板 ===")
    
    try:
        doc = Document(template_path)
        
        # 统计段落数和样式
        styles_used = {}
        headings = []
        
        for i, para in enumerate(doc.paragraphs):
            style = para.style.name if para.style else 'None'
            styles_used[style] = styles_used.get(style, 0) + 1
            
            if 'Heading' in style or '标题' in para.text[:10]:
                headings.append((i, para.text[:50], style))
        
        print(f"  总段落数: {len(doc.paragraphs)}")
        print(f"  表格数: {len(doc.tables)}")
        print(f"  使用的样式数: {len(styles_used)}")
        print(f"\n  主要样式:")
        for style, count in sorted(styles_used.items(), key=lambda x: -x[1])[:10]:
            print(f"    {style}: {count}")
        
        print(f"\n  标题位置 (前10个):")
        for idx, text, style in headings[:10]:
            print(f"    [{idx}] {style}: {text}")
        
        return len(doc.paragraphs)
        
    except Exception as e:
        print(f"  分析失败: {e}")
        return 0

def analyze_source(source_path):
    """Step 3: 分析源文档"""
    print("\n=== Step 3: 分析源文档 ===")
    
    try:
        doc = Document(source_path)
        
        styles_used = {}
        headings = []
        
        for i, para in enumerate(doc.paragraphs):
            style = para.style.name if para.style else 'None'
            styles_used[style] = styles_used.get(style, 0) + 1
            
            if 'Heading' in style:
                headings.append((i, para.text[:80], style))
        
        print(f"  总段落数: {len(doc.paragraphs)}")
        print(f"  表格数: {len(doc.tables)}")
        # 统计图片数
        img_count = 0
        for p in doc.paragraphs:
            for r in p.runs:
                if r._element.xpath('.//a:blip'):
                    img_count += 1
        print(f"  图片数: {img_count}")
        
        print(f"\n  主要样式:")
        for style, count in sorted(styles_used.items(), key=lambda x: -x[1])[:10]:
            print(f"    {style}: {count}")
        
        print(f"\n  章节标题:")
        for idx, text, style in headings[:15]:
            print(f"    [{idx}] {style}: {text}")
        
        return len(doc.paragraphs)
        
    except Exception as e:
        print(f"  分析失败: {e}")
        return 0

def apply_template_c1_overlay(source_path, output_path):
    """
    C-1: Overlay 方法
    保留模板样式，将源文档内容应用模板样式
    """
    print("\n=== 应用模板 (C-1: Overlay) ===")
    
    # 读取源文档内容
    source = Document(source_path)
    
    # 创建新文档
    doc = Document()
    
    # 复制源文档的每个段落
    for para in source.paragraphs:
        new_para = doc.add_paragraph()
        
        # 复制文本
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            # 保留基本格式
            if run.bold:
                new_run.bold = True
            if run.italic:
                new_run.italic = True
        
        # 复制样式名称（如果模板中有对应的样式）
        try:
            new_para.style = para.style
        except:
            pass
    
    # 复制表格
    for table in source.tables:
        new_table = doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text
    
    # 设置页面格式
    section = doc.sections[0]
    section.page_width = Cm(21)    # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    
    # 保存
    doc.save(output_path)
    print(f"  已保存: {output_path}")
    
    return True

def create_formatted_thesis(source_path, output_path):
    """
    创建符合成都大学规范的论文格式
    """
    print("\n=== 创建格式化论文 ===")
    
    source = Document(source_path)
    doc = Document()
    
    # 页面设置 (成都大学规范)
    section = doc.sections[0]
    section.page_width = Cm(21)    # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    
    # 处理每个段落
    processed_count = 0
    for para in source.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        new_para = doc.add_paragraph()
        
        # 判断段落类型并设置样式
        style_name = 'Normal'
        font_size = 12
        alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 标题识别
        if text.startswith(('第1章', '第2章', '第3章', '第4章', '第5章', '第6章', '第一章', '第二章')):
            style_name = 'Heading 1'
            font_size = 18
            alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif len(text) < 30 and any(marker in text for marker in ['1.', '1、', '（1）', '(1)']):
            style_name = 'Heading 2'
            font_size = 14
        elif text.startswith('【摘要】') or text == '摘要':
            style_name = 'Heading 2'
            font_size = 14
        elif text in ['Abstract', '参考文献', '致谢', '附录', '作者简介']:
            style_name = 'Heading 2'
            font_size = 14
        elif '图' in text[:5] and '示' in text[:10] and len(text) < 30:
            alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 应用样式
        try:
            new_para.style = style_name
        except:
            pass
        
        new_para.alignment = alignment
        
        # 添加文本
        run = new_para.add_run(text)
        run.font.size = Pt(font_size)
        
        # 中文书籍用宋体，英文用 Times New Roman
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        processed_count += 1
    
    # 复制表格
    for table in source.tables:
        new_table = doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text
                # 表格样式
                for para in new_table.rows[i].cells[j].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加表格后的空行
        doc.add_paragraph()
    
    # 保存
    doc.save(output_path)
    print(f"  已保存: {output_path}")
    print(f"  处理段落数: {processed_count}")
    
    return output_path

if __name__ == '__main__':
    print("=" * 60)
    print("Pipeline C: 成都大学论文模板套用")
    print("=" * 60)
    
    template, source, output_dir = check_files()
    
    if not os.path.exists(source):
        print(f"错误: 源文件 {source} 不存在!")
        exit(1)
    
    # 检查模板文件格式
    template_is_docx = False
    try:
        with zipfile.ZipFile(template, 'r') as z:
            template_is_docx = any('word/document.xml' in n for n in z.namelist())
    except:
        pass
    
    if not template_is_docx:
        # .doc 文件，需要转换
        converted_template = os.path.join(output_dir, 'template_converted.docx')
        if convert_doc_to_docx(template, converted_template):
            template = converted_template
            template_is_docx = True
        else:
            print("\n警告: .doc 转换失败，将使用基本格式设置")
    
    # 分析文档
    template_para_count = analyze_template(source) if template_is_docx else 0
    source_para_count = analyze_source(source)
    
    # 决定使用 C-1 还是 C-2
    # 由于模板是 .doc 格式，我们使用 C-1 Overlay 方法
    # 直接在源文档上应用标准论文格式
    
    output_path = os.path.join(output_dir, 'thesis_chengdu_university.docx')
    create_formatted_thesis(source, output_path)
    
    print("\n" + "=" * 60)
    print("完成!")
    print("=" * 60)
